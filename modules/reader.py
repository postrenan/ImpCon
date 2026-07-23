import re
from pathlib import Path

SUPPORTED = {".docx", ".pdf", ".txt", ".md"}


def read_document(file_path: str) -> str:
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext not in SUPPORTED:
        raise ValueError(f"Formato não suportado: '{ext}'. Use: {', '.join(SUPPORTED)}")
    if ext == ".docx":
        return _read_docx(path)
    elif ext == ".pdf":
        return _read_pdf(path)
    else:
        return path.read_text(encoding="utf-8", errors="replace")


def _read_docx(path: Path) -> str:
    from docx import Document
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("[Tabela] " + " | ".join(cells))
    return "\n\n".join(parts)


def _read_pdf(path: Path) -> str:
    """
    Use pdfplumber for spatial word grouping (avoids word-by-word extraction).
    Falls back to pypdf if pdfplumber is unavailable.
    """
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if text and text.strip():
                    pages.append(text.strip())
        raw = "\n\n".join(pages)
    except ImportError:
        raw = _read_pdf_pypdf(path)

    return _normalize(raw)


def _read_pdf_pypdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text and text.strip():
            pages.append(text.strip())
    return "\n\n".join(pages)


def _normalize(text: str) -> str:
    """
    Reconstruct paragraphs from line-by-line extracted text.
    pdfplumber gives one real line per \\n; we group them into paragraphs
    separated by blank lines, flushing at sentence-ending punctuation or
    when a section header is detected.
    """
    HEADER_RE = re.compile(
        r"^(CLÁUSULA|CLAUSULA|ARTIGO|ART\.?|CAPÍTULO|CAPITULO|SEÇÃO|SECAO)",
        re.IGNORECASE,
    )
    SENTENCE_END = re.compile(r"[.!?]\s*$")

    lines = re.sub(r" {2,}", " ", text).splitlines()

    paragraphs: list[str] = []
    current: list[str] = []

    def flush():
        if current:
            paragraphs.append(" ".join(current))
            current.clear()

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            flush()
            continue
        if HEADER_RE.match(line):
            flush()
            paragraphs.append(line)
            continue
        current.append(line)
        if SENTENCE_END.search(line):
            flush()

    flush()
    return "\n\n".join(p for p in paragraphs if p)
